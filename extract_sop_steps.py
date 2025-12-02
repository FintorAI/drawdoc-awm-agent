#!/usr/bin/env python3
"""
Extract Steps 1-28 from Docs Draw SOP document.

This script reads the DOCX file and extracts all steps, providing a clear summary
of each step's purpose and key actions.
"""

from docx import Document
from pathlib import Path
import re
import json
from typing import Dict, List, Any


def extract_steps_1_28(docx_path: Path) -> Dict[int, Dict[str, Any]]:
    """Extract and organize steps 1-28 from the SOP document.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary mapping step numbers to step information
    """
    doc = Document(docx_path)
    
    # Extract all paragraphs with text
    all_paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_paragraphs.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    all_paragraphs.append(text)
    
    steps = {}
    current_step = None
    current_content = []
    
    for i, text in enumerate(all_paragraphs):
        # Match various STEP patterns - prioritize main document steps
        # Patterns: STEP-1, STEP 1, STEP-1:, STEP 1:, STEP-1 –, etc.
        step_patterns = [
            r'^STEP[- ]?(\d+)[:–]\s*(.+)$',  # STEP-1: Title or STEP-1 – Title
            r'^STEP[- ]?(\d+)[:–]\s*$',  # STEP-1: or STEP-1 – (no title)
            r'^STEP[- ]?(\d+)\s+(.+)$',  # STEP 1 Title (with space)
            r'^STEP[- ]?(\d+)$',  # Just STEP-1
        ]
        
        matched = False
        for pattern in step_patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                step_num = int(match.group(1))
                step_title = match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else ""
                
                # Save previous step if exists
                if current_step and 1 <= current_step <= 28:
                    # Clean up content
                    clean_content = [c for c in current_content if c.strip() and len(c.strip()) > 5]
                    if clean_content:
                        # Extract title from header if not already extracted
                        header_text = clean_content[0]
                        if not step_title:
                            title_match = re.search(r'STEP[- ]?\d+[:–]?\s*(.+)', header_text, re.IGNORECASE)
                            if title_match:
                                step_title = title_match.group(1).strip()[:200]
                        
                        steps[current_step] = {
                            'full_header': header_text,
                            'title': step_title,
                            'content': clean_content[1:50] if len(clean_content) > 1 else [],  # First 50 content items
                            'content_count': len(clean_content) - 1 if len(clean_content) > 1 else 0
                        }
                
                # Start new step
                if 1 <= step_num <= 28:
                    current_step = step_num
                    current_content = [text] if text else []
                    matched = True
                    break
                else:
                    current_step = None
                    current_content = []
        
        if not matched and current_step and 1 <= current_step <= 28:
            # Continue collecting content for current step
            # Stop collecting if we hit a clear new major section
            is_new_section = False
            
            # Check for major section breaks
            if text and len(text) > 0:
                # Stop if it's a new STEP with different number
                other_step = re.match(r'^STEP[- ]?(\d+)', text, re.IGNORECASE)
                if other_step:
                    other_num = int(other_step.group(1))
                    if other_num != current_step and (other_num <= 28 or current_content):
                        is_new_section = True
                
                # Continue collecting unless it's clearly a new section
                if not is_new_section:
                    current_content.append(text)
    
    # Save last step
    if current_step and 1 <= current_step <= 28:
        clean_content = [c for c in current_content if c.strip() and len(c.strip()) > 5]
        if clean_content:
            step_title = ""
            header_text = clean_content[0]
            # Try to extract title from first line
            title_match = re.search(r'STEP[- ]?\d+[:–]?\s*(.+)', header_text, re.IGNORECASE)
            if title_match:
                step_title = title_match.group(1).strip()[:200]
            
            steps[current_step] = {
                'full_header': header_text,
                'title': step_title,
                'content': clean_content[1:50] if len(clean_content) > 1 else [],
                'content_count': len(clean_content) - 1 if len(clean_content) > 1 else 0
            }
    
    return steps


def print_summary(steps: Dict[int, Dict[str, Any]]):
    """Print a formatted summary of all steps."""
    print("="*80)
    print("STEPS 1-28 SUMMARY - DOCS DRAW SOP")
    print("="*80)
    print()
    
    for step_num in sorted(steps.keys()):
        if 1 <= step_num <= 28:
            step = steps[step_num]
            print(f"STEP {step_num}:")
            print("-" * 80)
            
            if step['title']:
                print(f"Title: {step['title']}")
                print()
            
            if step['content']:
                print("Key Points:")
                for i, content_item in enumerate(step['content'][:15], 1):  # Show first 15 items
                    # Clean up Encompass paths
                    clean_item = content_item.replace('Encompass >>', 'Encompass').strip()
                    # Truncate long lines
                    if len(clean_item) > 200:
                        clean_item = clean_item[:200] + "..."
                    print(f"  {i}. {clean_item}")
                
                if step['content_count'] > 15:
                    print(f"  ... ({step['content_count'] - 15} more items)")
            
            print()
            print()


def save_to_json(steps: Dict[int, Dict[str, Any]], output_path: Path):
    """Save steps to JSON file."""
    output_data = {
        "steps": steps,
        "total_steps": len(steps),
        "steps_found": sorted(steps.keys())
    }
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Steps saved to: {output_path}")


def main():
    """Main function."""
    # Find the DOCX file
    docx_path = Path("Docs Draw SOP (1).docx")
    
    if not docx_path.exists():
        # Try alternative locations
        alt_paths = [
            Path("agents/drawdocs/subagents/verification_agent/data/Docs Draw SOP (1).docx"),
            Path("../Docs Draw SOP (1).docx"),
        ]
        for alt_path in alt_paths:
            if alt_path.exists():
                docx_path = alt_path
                break
    
    if not docx_path.exists():
        print(f"Error: Could not find 'Docs Draw SOP (1).docx'")
        print(f"Searched in: {Path.cwd()}")
        return
    
    print(f"Reading: {docx_path}")
    print()
    
    # Extract steps
    try:
        steps = extract_steps_1_28(docx_path)
    except Exception as e:
        print(f"Error extracting steps: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # Print summary
    print_summary(steps)
    
    # Save to JSON
    output_json = Path("sop_steps_1_28.json")
    save_to_json(steps, output_json)
    
    print()
    print(f"Found {len(steps)} steps out of 28")
    missing = [i for i in range(1, 29) if i not in steps]
    if missing:
        print(f"Missing steps: {missing}")
    else:
        print("✓ All steps 1-28 found!")


if __name__ == "__main__":
    main()

