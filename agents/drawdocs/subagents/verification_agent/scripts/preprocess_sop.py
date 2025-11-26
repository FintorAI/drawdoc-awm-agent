"""
Preprocess SOP document to extract rules and create JSON index by page number.
"""

import json
import re
from pathlib import Path
from docx import Document
from typing import Dict, List, Any


def extract_sop_rules_from_docx(docx_path: Path) -> Dict[str, Any]:
    """
    Extract SOP rules from DOCX file and organize by page references.
    
    Args:
        docx_path: Path to the SOP DOCX file
        
    Returns:
        Dictionary mapping page numbers to rules and related fields
    """
    doc = Document(docx_path)
    
    # Extract all paragraphs with their text
    all_text = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            all_text.append(text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    all_text.append(text)
    
    # Organize content by sections/headers
    sop_rules = {}
    current_section = None
    current_content = []
    
    for text in all_text:
        # Check if this is a section header (all caps, or ends with colon, etc.)
        if text.isupper() or text.endswith(':'):
            if current_section and current_content:
                sop_rules[current_section] = ' '.join(current_content)
            current_section = text
            current_content = []
        else:
            current_content.append(text)
    
    # Add the last section
    if current_section and current_content:
        sop_rules[current_section] = ' '.join(current_content)
    
    # Create page-indexed structure
    # Since DOCX doesn't have reliable page numbers, we'll use sections
    page_indexed_rules = {
        "sections": sop_rules,
        "full_text": ' '.join(all_text)
    }
    
    return page_indexed_rules


def create_page_index_from_csv_references(
    csv_path: Path, 
    sop_content: Dict[str, Any]
) -> Dict[str, Dict[str, Any]]:
    """
    Create page-indexed rules based on CSV field mapping SOP page references.
    
    Args:
        csv_path: Path to the DrawingDoc Verifications CSV
        sop_content: Extracted SOP content
        
    Returns:
        Dictionary mapping page numbers to fields and rules
    """
    import csv
    
    page_index = {}
    
    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            field_name = row.get('Name', '')
            field_id = row.get('ID', '')
            sop_pages = row.get('SOP Pages', '')
            notes = row.get('Notes', '')
            
            if not sop_pages or not field_id:
                continue
            
            # Parse page numbers (can be comma-separated)
            pages = [p.strip() for p in sop_pages.split(',') if p.strip()]
            
            for page in pages:
                if page not in page_index:
                    page_index[page] = {
                        "fields": [],
                        "rules": []
                    }
                
                # Add field reference
                field_info = {
                    "field_id": field_id,
                    "field_name": field_name,
                    "notes": notes
                }
                
                page_index[page]["fields"].append(field_info)
                
                # Extract rules from notes
                if notes:
                    # Split notes by semicolon or period
                    rules = [r.strip() for r in re.split(r'[;.]', notes) if r.strip()]
                    page_index[page]["rules"].extend(rules)
    
    # Deduplicate rules per page
    for page in page_index:
        page_index[page]["rules"] = list(set(page_index[page]["rules"]))
    
    return page_index


def main():
    """Main function to preprocess SOP and create JSON output."""
    
    # Paths
    base_dir = Path(__file__).parent.parent
    docx_path = base_dir / "data" / "Docs Draw SOP (1).docx"
    csv_path = base_dir.parent.parent.parent.parent / "packages" / "data" / "DrawingDoc Verifications - Sheet1.csv"
    output_path = base_dir / "config" / "sop_rules.json"
    
    print(f"Extracting SOP from: {docx_path}")
    
    # Extract SOP content
    sop_content = extract_sop_rules_from_docx(docx_path)
    
    print(f"Creating page index from CSV: {csv_path}")
    
    # Create page-indexed rules based on CSV references
    page_indexed_rules = create_page_index_from_csv_references(csv_path, sop_content)
    
    # Combine both for comprehensive output
    output_data = {
        "page_indexed_rules": page_indexed_rules,
        "sop_sections": sop_content.get("sections", {}),
        "metadata": {
            "source_docx": str(docx_path.name),
            "source_csv": str(csv_path.name),
            "total_pages_referenced": len(page_indexed_rules)
        }
    }
    
    # Save to JSON
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    print(f"✓ SOP rules saved to: {output_path}")
    print(f"  - Total pages referenced: {len(page_indexed_rules)}")
    print(f"  - Total sections: {len(sop_content.get('sections', {}))}")
    
    # Print sample
    print("\nSample page index (Page 20):")
    if "20" in page_indexed_rules:
        print(json.dumps(page_indexed_rules["20"], indent=2)[:500])


if __name__ == "__main__":
    main()

